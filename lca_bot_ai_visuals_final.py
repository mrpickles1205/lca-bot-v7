
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import random
import os

def generate_lci_data():
    return pd.DataFrame({
        'Life Cycle Stage': ['Materials', 'Manufacturing', 'Use Phase', 'End-of-Life'],
        'Energy Use (MJ)': [random.uniform(80, 120), random.uniform(50, 100), random.uniform(10, 20), random.uniform(15, 30)],
        'GHG Emissions (kg CO2-eq)': [random.uniform(5, 10), random.uniform(8, 12), random.uniform(1, 3), random.uniform(2, 4)],
        'Water Use (L)': [random.uniform(20, 40), random.uniform(10, 30), random.uniform(1, 5), random.uniform(5, 15)]
    })

def create_visuals(df):
    chart_files = []

    for column in df.columns[1:]:
        # Bar chart
        fig, ax = plt.subplots()
        ax.bar(df['Life Cycle Stage'], df[column], color='steelblue')
        ax.set_title(f'{column} by Stage')
        bar_path = f'bar_{column}.png'
        fig.savefig(bar_path)
        chart_files.append(bar_path)
        plt.close(fig)

        # Pie chart
        fig, ax = plt.subplots()
        ax.pie(df[column], labels=df['Life Cycle Stage'], autopct='%1.1f%%', startangle=90)
        ax.set_title(f'{column} Distribution')
        pie_path = f'pie_{column}.png'
        fig.savefig(pie_path)
        chart_files.append(pie_path)
        plt.close(fig)

        # Line chart
        fig, ax = plt.subplots()
        ax.plot(df['Life Cycle Stage'], df[column], marker='o', linestyle='-', color='forestgreen')
        ax.set_title(f'{column} Trend')
        line_path = f'line_{column}.png'
        fig.savefig(line_path)
        chart_files.append(line_path)
        plt.close(fig)

    return chart_files

def ai_generated_content(product):
    return {
        "executive_summary": f"The LCA of the {product} reveals significant environmental impacts primarily during manufacturing and disposal. This report synthesizes known data and simulates a full cradle-to-grave analysis.",
        "recommendations": f"Recommendations for the {product} include modular design, recyclable batteries, and public awareness around e-waste.",
        "interpretation": f"The use phase has a minor cumulative effect, while manufacturing contributes the largest emissions due to energy-intensive materials like lithium batteries and plastics."
    }

def create_report(product, df, charts, ai_content):
    doc = Document()
    doc.add_heading(f'LCA Report for: {product}', 0)
    doc.add_paragraph(f'Date: {datetime.date.today()}')
    doc.add_paragraph("Confidential – For Internal Use Only").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc = [
        "Executive Summary", "1. Introduction", "2. Goal and Scope", "3. Functional Unit", "4. System Boundary",
        "5. Inventory Analysis", "6. Life Cycle Impact Assessment (LCIA)", "7. Interpretation",
        "8. Assumptions and Limitations", "9. Recommendations", "Appendix A: Glossary", "Appendix B: References"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Sections with AI content
    sections = {
        "Executive Summary": ai_content["executive_summary"],
        "1. Introduction": "This report follows ISO 14040 and 14044 standards to evaluate environmental impacts.",
        "2. Goal and Scope": "To assess the full life cycle environmental impact of a product from raw materials to disposal.",
        "3. Functional Unit": f"1 {product} used over an average 3-year lifespan.",
        "4. System Boundary": "Cradle-to-grave: includes raw materials, manufacturing, transport, use, and end-of-life.",
        "5. Inventory Analysis": "Detailed data of material and energy inputs/outputs collected and modeled.",
        "6. Life Cycle Impact Assessment (LCIA)": """Visual representation of environmental burdens by stage and impact type.""",
        "7. Interpretation": ai_content["interpretation"],
        "8. Assumptions and Limitations": "Assumptions include average usage rates, regional electricity mixes, and generalized transport models.",
        "9. Recommendations": ai_content["recommendations"],
        "Appendix A: Glossary": "LCA: Life Cycle Assessment
GWP: Global Warming Potential
MJ: Megajoules
CO2-eq: Carbon dioxide equivalent",
        "Appendix B: References": "1. ISO 14040:2006
2. ISO 14044:2006
3. ReCiPe 2016
4. IPCC AR6
5. Ecoinvent Database"
    }

    for title, text in sections.items():
        doc.add_heading(title, level=1 if not title.startswith("Appendix") else 2)
        doc.add_paragraph(text)
        if title == "6. Life Cycle Impact Assessment (LCIA)":
            for chart in charts:
                doc.add_paragraph(f"Figure: {chart.split('.')[0].replace('_', ' ').capitalize()}")
                doc.add_picture(chart, width=Inches(5.5))
        doc.add_page_break()

    # Add LCI Table
    doc.add_heading("Detailed LCI Table", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(round(val, 2)) if isinstance(val, (int, float)) else str(val)

    filename = f"LCA_Report_Visual_{product.replace(' ', '_')}.docx"
    doc.save(filename)
    return filename

# Streamlit app
st.title("🌿 ISO LCA Bot Pro")
product = st.text_input("Enter the product name:", "Electric Toothbrush")

if st.button("Generate LCA Report"):
    with st.spinner("Generating data, charts, and professional report..."):
        df = generate_lci_data()
        charts = create_visuals(df)
        ai_content = ai_generated_content(product)
        report_path = create_report(product, df, charts, ai_content)

    with open(report_path, "rb") as f:
        st.download_button("📄 Download DOCX Report", f, file_name=report_path)
